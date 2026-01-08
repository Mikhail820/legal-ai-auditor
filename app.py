import streamlit as st
import requests
import json
from PyPDF2 import PdfReader
from docx import Document
from docx.shared import Pt
from bs4 import BeautifulSoup
import io
import base64
import re
import time
from urllib.parse import urlparse
import logging

# Настройка логирования
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# -------------------
# 1. Настройки интерфейса
# -------------------
st.set_page_config(
    page_title="LegalAI Enterprise Pro", 
    page_icon="⚖️", 
    layout="wide",
    initial_sidebar_state="expanded"
)

# Кастомные стили
st.markdown("""
<style>
    /* Основные стили */
    .main-header { 
        font-size: 2.5rem; 
        color: #FF4B4B; 
        text-align: center; 
        margin-bottom: 1.5rem; 
        font-weight: 800;
        padding: 20px 0;
        background: linear-gradient(90deg, #FF4B4B 0%, #FF6B6B 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
    }
    
    .stButton>button { 
        width: 100%; 
        border-radius: 10px; 
        font-weight: bold; 
        height: 3.5em; 
        background: linear-gradient(135deg, #FF4B4B 0%, #FF6B6B 100%); 
        color: white; 
        border: none;
        transition: all 0.3s ease;
        margin-top: 10px;
    }
    
    .stButton>button:hover { 
        transform: translateY(-2px);
        box-shadow: 0 5px 15px rgba(255, 75, 75, 0.4);
    }
    
    .stDownloadButton>button { 
        width: 100%; 
        border-radius: 10px; 
        background: linear-gradient(135deg, #28a745 0%, #20c997 100%); 
        color: white; 
        border: none;
    }
    
    /* Карточки рисков */
    .risk-card { 
        background-color: #ffffff; 
        border-left: 6px solid #ff4b4b; 
        padding: 20px; 
        border-radius: 8px; 
        box-shadow: 0 4px 6px rgba(0,0,0,0.1); 
        margin-bottom: 20px; 
        color: #000;
        transition: transform 0.3s ease;
    }
    
    .risk-card:hover {
        transform: translateX(5px);
    }
    
    .score-container { 
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 25px; 
        border-radius: 15px; 
        text-align: center; 
        border: none;
        margin-bottom: 25px;
        color: white;
        box-shadow: 0 4px 15px rgba(102, 126, 234, 0.3);
    }
    
    .disclaimer { 
        font-size: 0.8rem; 
        color: #7f8c8d; 
        padding: 15px; 
        background: #fff3f3; 
        border-radius: 10px; 
        border: 1px solid #fab1a0;
        margin: 10px 0;
    }
    
    /* Вкладки */
    .stTabs [data-baseweb="tab-list"] {
        gap: 8px;
    }
    
    .stTabs [data-baseweb="tab"] {
        border-radius: 8px 8px 0 0;
        padding: 10px 20px;
        font-weight: 600;
    }
    
    /* Прогресс бар */
    .stProgress > div > div > div {
        background: linear-gradient(90deg, #FF4B4B 0%, #FF6B6B 100%);
    }
    
    /* Боковая панель */
    section[data-testid="stSidebar"] {
        background: linear-gradient(180deg, #2d3436 0%, #1a1e1f 100%);
    }
    
    section[data-testid="stSidebar"] h1, 
    section[data-testid="stSidebar"] h2, 
    section[data-testid="stSidebar"] h3,
    section[data-testid="stSidebar"] label {
        color: white !important;
    }
    
    /* Улучшенные текстовые области */
    .stTextArea textarea {
        border-radius: 10px;
        border: 2px solid #e0e0e0;
    }
    
    /* Карточки загрузки файлов */
    .upload-card {
        border: 2px dashed #667eea;
        border-radius: 10px;
        padding: 20px;
        text-align: center;
        background: #f8f9fa;
        margin: 10px 0;
    }
</style>
""", unsafe_allow_html=True)

# Текст дисклеймера
DISCLAIMER_TEXT = """
⚠️ **ВНИМАНИЕ:** 
Анализ выполнен искусственным интеллектом. Не является юридической консультацией. 
Все выводы требуют обязательной проверки у квалифицированного юриста. 
Используйте на свой страх и риск.
"""

# -------------------
# 2. Конфигурация моделей и API
# -------------------
MODEL_POLICY = [
    "gemini-2.0-flash-lite",
    "gemini-2.0-flash",
    "gemini-1.5-flash",
    "gemini-1.5-pro"
]

# Получение API ключа
try:
    API_KEY = st.secrets.get("GOOGLE_API_KEY")
    if not API_KEY:
        st.error("❌ API ключ не найден в Secrets. Добавьте GOOGLE_API_KEY.")
except:
    API_KEY = None
    st.warning("⚠️ Secrets не настроены. Укажите API ключ вручную.")

# Резервный ввод API ключа
if not API_KEY:
    with st.sidebar:
        API_KEY = st.text_input("🔑 Введите Google API Key:", type="password")
        if API_KEY:
            st.success("✅ Ключ принят")
        else:
            st.warning("Введите API ключ для работы приложения")

# -------------------
# 3. Улучшенная функция вызова Gemini
# -------------------
@st.cache_data(show_spinner=False, max_entries=10)
def call_gemini_safe(prompt: str, content: str, is_image: bool = False, model_override: str = None):
    """
    Безопасный вызов Gemini API с обработкой ошибок и ретраями
    """
    if not API_KEY:
        return "❌ Ошибка: Отсутствует API ключ. Добавьте GOOGLE_API_KEY в Secrets или введите вручную."
    
    if not content or (isinstance(content, str) and not content.strip()):
        return "⚠️ Предупреждение: Пустой документ или текст для анализа."
    
    # Ограничение длины текста для предотвращения перегрузки
    if isinstance(content, str) and len(content) > 100000:
        content = content[:100000] + "\n\n... [текст обрезан из-за большого объема]"
    
    models_to_try = [model_override] if model_override else MODEL_POLICY
    
    for model_idx, model in enumerate(models_to_try):
        try:
            url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={API_KEY}"
            
            # Формирование payload в зависимости от типа контента
            if is_image:
                # Определяем MIME тип по первым байтам
                if content[:3] == b'\xff\xd8\xff':
                    mime_type = "image/jpeg"
                elif content[:8] == b'\x89PNG\r\n\x1a\n':
                    mime_type = "image/png"
                else:
                    mime_type = "image/jpeg"  # fallback
                
                img_b64 = base64.b64encode(content).decode('utf-8')
                payload = {
                    "contents": [{
                        "parts": [
                            {"text": prompt},
                            {"inline_data": {"mime_type": mime_type, "data": img_b64}}
                        ]
                    }],
                    "generationConfig": {
                        "temperature": 0.1,
                        "topP": 0.8,
                        "topK": 40
                    }
                }
            else:
                payload = {
                    "contents": [{
                        "parts": [{"text": f"{prompt}\n\n=== ДОКУМЕНТ ДЛЯ АНАЛИЗА ===\n{content}\n=== КОНЕЦ ДОКУМЕНТА ==="}]
                    }],
                    "generationConfig": {
                        "temperature": 0.2,
                        "topP": 0.9,
                        "topK": 50,
                        "maxOutputTokens": 4000
                    }
                }
            
            # Логирование (безопасное)
            logger.info(f"Вызов модели: {model}, длина контента: {len(content) if isinstance(content, str) else 'image'}")
            
            # Вызов API с таймаутом
            headers = {"Content-Type": "application/json"}
            response = requests.post(url, json=payload, headers=headers, timeout=120)
            
            if response.status_code == 200:
                result = response.json()
                if 'candidates' in result and result['candidates']:
                    text = result['candidates'][0]['content']['parts'][0]['text']
                    logger.info(f"Успешный ответ от модели {model}")
                    return text
                else:
                    logger.warning(f"Пустой ответ от модели {model}: {result}")
                    continue
            
            elif response.status_code == 429:
                logger.warning(f"Rate limit для модели {model}. Пробуем следующую...")
                time.sleep(1)  # Небольшая задержка перед ретраем
                continue
                
            else:
                error_msg = response.json().get('error', {}).get('message', 'Неизвестная ошибка')
                logger.error(f"Ошибка API ({model}): {response.status_code} - {error_msg}")
                continue
                
        except requests.exceptions.Timeout:
            logger.warning(f"Таймаут для модели {model}")
            continue
        except requests.exceptions.RequestException as e:
            logger.error(f"Сетевая ошибка для модели {model}: {str(e)}")
            continue
        except Exception as e:
            logger.error(f"Неожиданная ошибка для модели {model}: {str(e)}")
            continue
    
    return "⚠️ Все модели временно недоступны. Пожалуйста, попробуйте позже или проверьте:\n1. Доступность API ключа\n2. Интернет-соединение\n3. Лимиты API"

# -------------------
# 4. Инструменты для работы с документами
# -------------------
def extract_text(file_bytes: bytes, filename: str) -> str:
    """
    Извлечение текста из различных форматов файлов
    """
    try:
        filename_lower = filename.lower()
        
        if filename_lower.endswith(".pdf"):
            text_parts = []
            try:
                pdf_reader = PdfReader(io.BytesIO(file_bytes))
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        text = page.extract_text()
                        if text and text.strip():
                            text_parts.append(f"[Страница {page_num + 1}]\n{text}\n")
                    except Exception as e:
                        logger.warning(f"Ошибка чтения страницы {page_num}: {str(e)}")
                        continue
                
                if not text_parts:
                    return "⚠️ Не удалось извлечь текст из PDF. Возможно, документ состоит из сканированных изображений."
                    
                return "\n".join(text_parts)
                
            except Exception as e:
                return f"❌ Ошибка чтения PDF: {str(e)}"
        
        elif filename_lower.endswith(".docx"):
            try:
                doc = Document(io.BytesIO(file_bytes))
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                return "\n".join(paragraphs)
            except Exception as e:
                return f"❌ Ошибка чтения DOCX: {str(e)}"
        
        elif filename_lower.endswith(".txt"):
            try:
                return file_bytes.decode('utf-8', errors='ignore')
            except:
                return file_bytes.decode('cp1251', errors='ignore')
        
        else:
            return "❌ Неподдерживаемый формат файла. Используйте PDF, DOCX или TXT."
            
    except Exception as e:
        logger.error(f"Критическая ошибка при извлечении текста: {str(e)}")
        return f"❌ Не удалось прочитать файл: {str(e)}"

def validate_url(url: str) -> bool:
    """Проверка валидности URL"""
    try:
        result = urlparse(url)
        return all([result.scheme in ['http', 'https'], result.netloc])
    except:
        return False

def fetch_url_content(url: str) -> str:
    """Безопасное получение контента с веб-страницы"""
    try:
        if not validate_url(url):
            return "❌ Неверный URL. Используйте формат http:// или https://"
        
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        }
        
        response = requests.get(
            url, 
            headers=headers, 
            timeout=15,
            verify=True,
            allow_redirects=True
        )
        
        response.raise_for_status()
        
        # Проверка типа контента
        content_type = response.headers.get('content-type', '').lower()
        if 'text/html' not in content_type:
            return f"⚠️ URL не содержит HTML. Content-Type: {content_type}"
        
        soup = BeautifulSoup(response.text, 'html.parser')
        
        # Удаляем скрипты и стили
        for script in soup(["script", "style", "nav", "footer", "header"]):
            script.decompose()
        
        # Извлекаем основной текст
        text = soup.get_text(separator='\n', strip=True)
        
        # Очистка лишних пробелов и переносов
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = '\n'.join(chunk for chunk in chunks if chunk)
        
        if not text:
            return "⚠️ Не удалось извлечь текст со страницы"
        
        # Ограничение длины
        if len(text) > 30000:
            text = text[:30000] + "\n\n... [контент обрезан]"
        
        return text
        
    except requests.exceptions.Timeout:
        return "❌ Таймаут при загрузке URL. Проверьте доступность сайта."
    except requests.exceptions.RequestException as e:
        return f"❌ Ошибка сети: {str(e)}"
    except Exception as e:
        return f"❌ Неожиданная ошибка: {str(e)}"

def create_docx(text: str, title: str) -> io.BytesIO:
    """
    Создание DOCX документа из текста с форматированием
    """
    try:
        doc = Document()
        
        # Заголовок
        title_para = doc.add_heading(title, 0)
        title_para.alignment = 1  # По центру
        
        # Дата
        from datetime import datetime
        date_str = datetime.now().strftime("%d.%m.%Y %H:%M")
        doc.add_paragraph(f"Дата создания: {date_str}").italic = True
        
        # Дисклеймер
        disclaimer_para = doc.add_paragraph(DISCLAIMER_TEXT)
        disclaimer_para.italic = True
        for run in disclaimer_para.runs:
            run.font.color.rgb = 0xFF0000  # Красный цвет
        
        doc.add_paragraph().add_run().add_break()  # Разделитель
        
        # Обработка текста
        lines = text.split('\n')
        table_data = []
        in_table = False
        
        for line in lines:
            line = line.rstrip()
            
            # Детекция таблицы Markdown
            if '|' in line and not re.match(r'^[\|\s\-:]+$', line.strip()):
                cells = [cell.strip() for cell in line.split('|') if cell.strip()]
                if cells:
                    table_data.append(cells)
                    in_table = True
                continue
            
            # Если собрана таблица, вставляем её
            if in_table and table_data and (not line.strip() or '|' not in line):
                if len(table_data) > 1:  # Минимум 2 строки для таблицы
                    # Определяем количество колонок
                    num_cols = max(len(row) for row in table_data)
                    table = doc.add_table(rows=len(table_data), cols=num_cols)
                    table.style = 'Table Grid'
                    
                    # Заполняем таблицу
                    for i, row in enumerate(table_data):
                        for j, cell_text in enumerate(row):
                            if j < num_cols:
                                cell = table.cell(i, j)
                                cell.text = cell_text
                                # Центрируем заголовки
                                if i == 0:
                                    for paragraph in cell.paragraphs:
                                        paragraph.alignment = 1
                
                table_data = []
                in_table = False
                doc.add_paragraph()  # Отступ после таблицы
            
            # Обработка обычного текста
            if line.strip() and not in_table:
                # Убираем Markdown разметку
                clean_line = re.sub(r'^[#\*\-\+]+|\*\*|\*|__|_|~~', '', line).strip()
                
                if line.startswith('## '):
                    doc.add_heading(clean_line, 2)
                elif line.startswith('# '):
                    doc.add_heading(clean_line, 1)
                elif line.startswith('### '):
                    doc.add_heading(clean_line, 3)
                elif line.startswith('- ') or line.startswith('* ') or line.startswith('+ '):
                    doc.add_paragraph(clean_line, style='List Bullet')
                elif re.match(r'^\d+\.', line):
                    doc.add_paragraph(clean_line, style='List Number')
                else:
                    para = doc.add_paragraph(clean_line)
                    
                    # Выделение ключевых фраз
                    if any(keyword in line for keyword in ['риск', 'опасность', 'проблема', '⚠️', '🔴']):
                        for run in para.runs:
                            run.bold = True
                            run.font.color.rgb = 0xFF0000
                    elif any(keyword in line for keyword in ['рекомендация', 'совет', 'решение', '✅', '💡']):
                        for run in para.runs:
                            run.font.color.rgb = 0x008000
        
        # Сохранение в буфер
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        
        logger.info(f"DOCX создан успешно: {title}, размер: {len(buf.getvalue())} байт")
        return buf
        
    except Exception as e:
        logger.error(f"Ошибка создания DOCX: {str(e)}")
        # Возвращаем простой текстовый файл в случае ошибки
        buf = io.BytesIO()
        buf.write(f"Ошибка создания документа: {str(e)}\n\n{text}".encode('utf-8'))
        buf.seek(0)
        return buf

# -------------------
# 5. Боковая панель
# -------------------
with st.sidebar:
    st.markdown("### ⚙️ Конфигурация анализа")
    
    role = st.radio(
        "👤 Анализ для:",
        ["Предприниматель", "Юрист", "Физическое лицо", "Корпоративный клиент"],
        help="Настройка рекомендаций под вашу роль"
    )
    
    loc = st.selectbox(
        "🌍 Юрисдикция:",
        ["Российская Федерация", "Казахстан", "Узбекистан", "Беларусь", "Международное право"],
        index=0
    )
    
    detail = st.select_slider(
        "📊 Уровень детализации:",
        options=["Краткий обзор", "Стандартный", "Детальный анализ", "Максимальный"],
        value="Стандартный",
        help="Влияет на объем и глубину анализа"
    )
    
    st.divider()
    
    st.markdown("### 🎯 Дополнительные параметры")
    
    include_recommendations = st.checkbox("Включить рекомендации", value=True)
    include_alternatives = st.checkbox("Показать альтернативные формулировки", value=False)
    
    st.divider()
    
    # Кэш и сброс
    col_cache1, col_cache2 = st.columns(2)
    with col_cache1:
        if st.button("🗑️ Очистить кэш", use_container_width=True):
            st.cache_data.clear()
            st.success("Кэш очищен!")
            time.sleep(1)
            st.rerun()
    
    with col_cache2:
        if st.button("🔄 Сбросить всё", use_container_width=True):
            for key in list(st.session_state.keys()):
                del st.session_state[key]
            st.success("Все данные сброшены!")
            time.sleep(1)
            st.rerun()
    
    st.divider()
    
    # Дисклеймер
    with st.expander("⚠️ Важная информация", expanded=True):
        st.markdown(f'<div class="disclaimer">{DISCLAIMER_TEXT}</div>', unsafe_allow_html=True)
    
    # Информация о системе
    st.markdown("---")
    st.markdown("**LegalAI Enterprise Pro v2.0**")
    st.caption("Анализ документов с использованием Google Gemini AI")

# -------------------
# 6. Основной интерфейс
# -------------------
st.markdown('<div class="main-header">⚖️ LegalAI Enterprise Pro</div>', unsafe_allow_html=True)

# Инициализация сессионных переменных
if 'audit_result' not in st.session_state:
    st.session_state.audit_result = None
if 'comparison_result' not in st.session_state:
    st.session_state.comparison_result = None

# Вкладки
tab1, tab2, tab3 = st.tabs(["🚀 УМНЫЙ АУДИТ ДОКУМЕНТОВ", "🔍 СРАВНЕНИЕ ВЕРСИЙ", "📋 ГЕНЕРАТОР ДОКУМЕНТОВ"])

# -------------------
# ВКЛАДКА 1: Умный аудит
# -------------------
with tab1:
    st.markdown("### Анализ рисков и юридической корректности документов")
    
    col1, col2 = st.columns([1, 1.2], gap="large")
    
    with col1:
        st.markdown("#### 📄 Параметры анализа")
        
        # Выбор типа документа
        doc_type = st.selectbox(
            "Тип документа:",
            [
                "Договор оказания услуг",
                "Договор поставки",
                "Договор аренды",
                "Соглашение о конфиденциальности (NDA)",
                "Трудовой договор",
                "Договор подряда",
                "Лицензионное соглашение",
                "Договор купли-продажи",
                "Агентский договор",
                "Другое"
            ],
            index=0,
            help="Выберите наиболее подходящий тип документа для более точного анализа"
        )
        
        # Способ ввода
        input_method = st.radio(
            "Способ ввода:",
            ["Загрузка файла", "Вставка текста", "URL веб-страницы"],
            horizontal=True,
            help="Выберите способ загрузки документа для анализа"
        )
        
        input_data = None
        is_image_file = False
        
        # Обработка разных способов ввода
        if input_method == "Загрузка файла":
            uploaded_file = st.file_uploader(
                "Загрузите документ (PDF, DOCX, TXT) или изображение (JPG, PNG):",
                type=["pdf", "docx", "txt", "png", "jpg", "jpeg"],
                help="Максимальный размер файла: 50MB"
            )
            
            if uploaded_file:
                # Проверка размера файла
                if uploaded_file.size > 50 * 1024 * 1024:  # 50MB
                    st.error("❌ Файл слишком большой. Максимальный размер: 50MB")
                else:
                    # Определяем тип файла
                    if uploaded_file.type.startswith("image"):
                        st.image(uploaded_file, caption="Загруженное изображение", use_column_width=True)
                        input_data = uploaded_file.getvalue()
                        is_image_file = True
                        st.success(f"✅ Изображение загружено: {uploaded_file.name}")
                    else:
                        # Извлекаем текст из документа
                        with st.spinner("Извлекаю текст из документа..."):
                            input_data = extract_text(uploaded_file.getvalue(), uploaded_file.name)
                        
                        if input_data.startswith("❌") or input_data.startswith("⚠️"):
                            st.error(input_data)
                        else:
                            st.success(f"✅ Текст извлечен ({len(input_data)} символов)")
                            with st.expander("📝 Предпросмотр текста"):
                                st.text_area("Извлеченный текст:", input_data[:2000] + ("..." if len(input_data) > 2000 else ""), height=200)
        
        elif input_method == "Вставка текста":
            input_data = st.text_area(
                "Вставьте текст документа:",
                height=300,
                placeholder="Вставьте сюда текст договора или другого документа для анализа...",
                help="Вы можете скопировать текст из любого документа и вставить его здесь"
            )
            
            if input_data:
                st.info(f"📝 Длина текста: {len(input_data)} символов")
        
        elif input_method == "URL веб-страницы":
            url_input = st.text_input(
                "Введите URL документа:",
                placeholder="https://example.com/document.html",
                help="Введите корректный URL веб-страницы, содержащей текст документа"
            )
            
            if url_input:
                with st.spinner("Загружаю и обрабатываю веб-страницу..."):
                    input_data = fetch_url_content(url_input)
                
                if input_data.startswith("❌") or input_data.startswith("⚠️"):
                    st.error(input_data)
                else:
                    st.success(f"✅ Контент загружен ({len(input_data)} символов)")
                    with st.expander("📝 Предпросмотр контента"):
                        st.text_area("Извлеченный текст:", input_data[:2000] + ("..." if len(input_data) > 2000 else ""), height=200)
        
        # Кнопка запуска анализа
        analyze_button = st.button(
            "🚀 ЗАПУСТИТЬ АНАЛИЗ РИСКОВ",
            type="primary",
            disabled=not input_data or not API_KEY,
            use_container_width=True
        )
    
    with col2:
        st.markdown("#### 📊 Результаты анализа")
        
        if analyze_button and input_data:
            with st.spinner("🤖 Анализирую документ..."):
                # Создаем прогресс бар
                progress_bar = st.progress(0)
                
                # Формируем промпт для анализа
                prompt = f"""
                Ты - опытный юрист-аналитик. Проведи экспертный анализ документа.
                
                КОНТЕКСТ:
                - Роль пользователя: {role}
                - Юрисдикция: {loc}
                - Тип документа: {doc_type}
                - Уровень детализации: {detail}
                
                ТРЕБОВАНИЯ К АНАЛИЗУ:
                1. Определи LEGAL SCORE (оценка юридической корректности) от 0% до 100%
                2. Выдели ключевые риски с пометкой 🔴
                3. Укажи возможные финансовые потери с пометкой 💸
                4. Отметь скрытые ловушки и неявные условия с пометкой ⚠️
                5. Проверь соответствие законодательству {loc}
                6. Проанализируй баланс сторон
                7. Оцени clarity (ясность формулировок)
                
                {"8. Предложи рекомендации по улучшению" if include_recommendations else ""}
                {"9. Приведи альтернативные формулировки критических пунктов" if include_alternatives else ""}
                
                ФОРМАТ ВЫВОДА:
                ## 📊 LEGAL SCORE: X%
                [Краткое резюме оценки]
                
                ## 🔴 КЛЮЧЕВЫЕ РИСКИ
                - Риск 1: [описание, уровень опасности, последствия]
                - Риск 2: [описание, уровень опасности, последствия]
                
                ## 💸 ФИНАНСОВЫЕ АСПЕКТЫ
                - [Потенциальные убытки, штрафы, издержки]
                
                ## ⚠️ СКРЫТЫЕ ЛОВУШКИ
                - [Проблемные формулировки, двусмысленности]
                
                {"## 💡 РЕКОМЕНДАЦИИ" if include_recommendations else ""}
                {"- [Конкретные предложения по улучшению]" if include_recommendations else ""}
                
                {"## 🔄 АЛЬТЕРНАТИВНЫЕ ФОРМУЛИРОВКИ" if include_alternatives else ""}
                {"| Пункт | Текущая формулировка | Рекомендуемая формулировка | Обоснование |" if include_alternatives else ""}
                {"|---|---|---|---|" if include_alternatives else ""}
                
                Будь конкретным, цитируй проблемные места из документа.
                """
                
                progress_bar.progress(30)
                
                # Вызов модели
                analysis_result = call_gemini_safe(
                    prompt, 
                    input_data, 
                    is_image_file,
                    model_override="gemini-2.0-flash"
                )
                
                progress_bar.progress(80)
                
                if analysis_result:
                    st.session_state.audit_result = analysis_result
                    progress_bar.progress(100)
                    time.sleep(0.5)
                    progress_bar.empty()
                    
                    st.success("✅ Анализ завершен!")
                else:
                    st.error("❌ Не удалось выполнить анализ. Попробуйте еще раз.")
                    progress_bar.empty()
        
        # Отображение результатов анализа
        if st.session_state.audit_result:
            st.markdown('<div class="score-container"><h3>📊 РЕЗУЛЬТАТЫ АНАЛИЗА</h3></div>', unsafe_allow_html=True)
            
            # Разбиваем результат на части для форматированного отображения
            result_lines = st.session_state.audit_result.split('\n')
            
            for line in result_lines:
                line_stripped = line.strip()
                
                # Выделение оценки
                if "LEGAL SCORE:" in line.upper() or "ОЦЕНКА:" in line.upper():
                    st.markdown(f"### {line}")
                    st.divider()
                
                # Карточки рисков
                elif any(marker in line for marker in ["🔴", "💸", "⚠️", "💡", "🔄"]):
                    st.markdown(f'<div class="risk-card">{line}</div>', unsafe_allow_html=True)
                
                # Обычный текст
                elif line_stripped:
                    st.markdown(line)
            
            # Кнопка скачивания
            st.divider()
            docx_file = create_docx(st.session_state.audit_result, f"Анализ документа: {doc_type}")
            
            col_dl1, col_dl2 = st.columns(2)
            with col_dl1:
                st.download_button(
                    label="📥 Скачать отчет (DOCX)",
                    data=docx_file,
                    file_name=f"Legal_Analysis_{doc_type}_{time.strftime('%Y%m%d_%H%M%S')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            
            with col_dl2:
                if st.button("📋 Копировать в буфер", use_container_width=True):
                    st.session_state.clipboard = st.session_state.audit_result
                    st.success("Текст скопирован в буфер обмена!")
        
        elif not analyze_button:
            st.info("👈 Загрузите документ и нажмите кнопку для запуска анализа")
            st.markdown("""
            ### Что анализирует система:
            - 📝 **Юридическую корректность** формулировок
            - ⚖️ **Соответствие законодательству** выбранной юрисдикции
            - 🔍 **Скрытые риски** и неявные условия
            - 💰 **Финансовые аспекты** и потенциальные убытки
            - ⚠️ **Ловушки** в договорной документации
            - ✅ **Баланс интересов** сторон договора
            """)

# -------------------
# ВКЛАДКА 2: Сравнение версий
# -------------------
with tab2:
    st.markdown("### Сравнение разных версий документа")
    st.info("Загрузите две версии одного документа для анализа изменений")
    
    col_a, col_b = st.columns(2, gap="large")
    
    with col_a:
        st.markdown("#### 📄 Оригинальная версия")
        file_a = st.file_uploader(
            "Загрузите оригинальный документ",
            type=["pdf", "docx", "txt"],
            key="file_a"
        )
        
        if file_a:
            st.success(f"✅ Загружено: {file_a.name}")
            text_a = extract_text(file_a.getvalue(), file_a.name)
            if not text_a.startswith("❌"):
                st.caption(f"Извлечено символов: {len(text_a)}")
    
    with col_b:
        st.markdown("#### 📄 Редактированная версия")
        file_b = st.file_uploader(
            "Загрузите измененную версию",
            type=["pdf", "docx", "txt"],
            key="file_b"
        )
        
        if file_b:
            st.success(f"✅ Загружено: {file_b.name}")
            text_b = extract_text(file_b.getvalue(), file_b.name)
            if not text_b.startswith("❌"):
                st.caption(f"Извлечено символов: {len(text_b)}")
    
    # Кнопка сравнения
    compare_button = st.button(
        "⚖️ СРАВНИТЬ ВЕРСИИ",
        type="primary",
        disabled=not (file_a and file_b) or not API_KEY,
        use_container_width=True
    )
    
    if compare_button and file_a and file_b:
        with st.spinner("🔍 Сравниваю документы..."):
            # Извлекаем тексты
            text_a = extract_text(file_a.getvalue(), file_a.name)
            text_b = extract_text(file_b.getvalue(), file_b.name)
            
            if text_a.startswith("❌") or text_b.startswith("❌"):
                st.error("Ошибка при чтении файлов. Проверьте формат документов.")
            else:
                # Промпт для сравнения
                compare_prompt = f"""
                Ты - юрист, специализирующийся на сравнении документов. Сравни две версии документа.
                
                ТРЕБОВАНИЯ:
                1. Создай таблицу сравнения в формате Markdown
                2. Для каждого значимого изменения укажи:
                   - Тип изменения (добавлено/удалено/изменено)
                   - Смысл изменения
                   - Юридические последствия
                   - Уровень риска (низкий/средний/высокий)
                
                3. Выдели критические изменения, влияющие на:
                   - Права сторон
                   - Обязательства
                   - Ответственность
                   - Финансовые условия
                
                4. В конце дай общую оценку:
                   - Насколько изменения улучшили/ухудшили документ
                   - Кому изменения выгодны
                   - Рекомендации по принятию/отклонению изменений
                
                ФОРМАТ ТАБЛИЦЫ:
                | Пункт | Было | Стало | Тип изменения | Риск | Комментарий |
                |-------|------|-------|---------------|------|-------------|
                
                ДОКУМЕНТ А (ОРИГИНАЛ):
                {text_a[:15000]}
                
                ДОКУМЕНТ Б (РЕДАКЦИЯ):
                {text_b[:15000]}
                """
                
                comparison_result = call_gemini_safe(compare_prompt, f"Сравнение {file_a.name} и {file_b.name}")
                
                if comparison_result:
                    st.session_state.comparison_result = comparison_result
                    st.success("✅ Сравнение завершено!")
                else:
                    st.error("❌ Не удалось сравнить документы")
    
    # Отображение результатов сравнения
    if st.session_state.comparison_result:
        st.markdown("### 📊 Результаты сравнения")
        st.markdown(st.session_state.comparison_result)
        
        # Кнопка скачивания
        docx_file = create_docx(
            st.session_state.comparison_result, 
            f"Сравнение документов: {file_a.name if 'file_a' in locals() else ''} vs {file_b.name if 'file_b' in locals() else ''}"
        )
        
        st.download_button(
            label="📥 Скачать отчет сравнения (DOCX)",
            data=docx_file,
            file_name=f"Document_Comparison_{time.strftime('%Y%m%d_%H%M%S')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

# -------------------
# ВКЛАДКА 3: Генератор документов
# -------------------
with tab3:
    st.markdown("### Генерация юридических документов")
    
    # Два режима генерации
    mode = st.radio(
        "Режим генерации:",
        ["На основе анализа", "С нуля по описанию"],
        horizontal=True
    )
    
    if mode == "На основе анализа":
        if st.session_state.audit_result:
            st.info("Сгенерируйте документы на основе проведенного анализа")
            
            col_gen1, col_gen2, col_gen3 = st.columns(3)
            
            with col_gen1:
                if st.button("📋 Протокол разногласий", use_container_width=True):
                    with st.spinner("Генерирую протокол разногласий..."):
                        prompt = """
                        На основе предыдущего анализа создай Протокол разногласий.
                        
                        Требования:
                        1. Формат - официальный документ
                        2. Структура:
                           - Шапка с реквизитами сторон
                           - Таблица разногласий
                           - Обоснования изменений
                           - Порядок урегулирования
                        
                        3. Таблица должна содержать:
                           | № п/п | Пункт договора | Редакция контрагента | Предлагаемая редакция | Обоснование |
                           |-------|----------------|----------------------|-----------------------|-------------|
                        
                        4. Учитывай риски, выявленные в анализе
                        5. Предложи юридически корректные формулировки
                        6. Укажи сроки рассмотрения
                        
                        Будь конкретным и практичным.
                        """
                        
                        protocol = call_gemini_safe(prompt, st.session_state.audit_result)
                        if protocol:
                            st.markdown(protocol)
                            
                            # Кнопка скачивания
                            docx_file = create_docx(protocol, "Протокол разногласий")
                            st.download_button(
                                "📥 Скачать Протокол",
                                data=docx_file,
                                file_name="Protocol_of_Disagreements.docx",
                                use_container_width=True
                            )
            
            with col_gen2:
                if st.button("✍️ Дополнительное соглашение", use_container_width=True):
                    with st.spinner("Генерирую дополнительное соглашение..."):
                        prompt = """
                        На основе анализа создай проект Дополнительного соглашения.
                        
                        Требования:
                        1. Официальная форма договора
                        2. Включи:
                           - Преамбулу
                           - Предмет соглашения
                           - Изменяемые условия
                           - Порядок вступления в силу
                           - Реквизиты сторон
                        
                        3. Конкретные формулировки изменений
                        4. Ссылки на пункты оригинального договора
                        5. Юридически корректный язык
                        
                        Сделай документ готовым к подписанию.
                        """
                        
                        agreement = call_gemini_safe(prompt, st.session_state.audit_result)
                        if agreement:
                            st.markdown(agreement)
                            
                            docx_file = create_docx(agreement, "Дополнительное соглашение")
                            st.download_button(
                                "📥 Скачать Соглашение",
                                data=docx_file,
                                file_name="Additional_Agreement.docx",
                                use_container_width=True
                            )
            
            with col_gen3:
                if st.button("📝 Правки для контрагента", use_container_width=True):
                    with st.spinner("Готовлю правки..."):
                        prompt = """
                        На основе анализа подготовь письмо контрагенту с предложением правок.
                        
                        Требования:
                        1. Деловой стиль переписки
                        2. Вежливый, но настойчивый тон
                        3. Конкретные предложения по изменению
                        4. Обоснование каждой правки
                        5. Сроки на рассмотрение
                        6. Контакты для обсуждения
                        
                        Сделай письмо убедительным и профессиональным.
                        """
                        
                        letter = call_gemini_safe(prompt, st.session_state.audit_result)
                        if letter:
                            st.markdown(letter)
                            
                            docx_file = create_docx(letter, "Письмо контрагенту")
                            st.download_button(
                                "📥 Скачать Письмо",
                                data=docx_file,
                                file_name="Letter_to_Counterparty.docx",
                                use_container_width=True
                            )
        else:
            st.warning("⚠️ Сначала выполните анализ документа на вкладке 'Умный аудит'")
    
    else:  # Режим "С нуля по описанию"
        st.markdown("#### 📝 Описание документа")
        
        doc_description = st.text_area(
            "Опишите, какой документ нужно создать:",
            height=150,
            placeholder="Например: 'Договор аренды офиса в Москве на 2 года с возможностью пролонгации. Арендодатель - юридическое лицо, арендатор - ИП. Гарантийный депозит 2 месяца. Ответственность за коммунальные платежи на арендаторе.'",
            help="Чем подробнее описание, тем точнее будет сгенерирован документ"
        )
        
        if doc_description:
            # Дополнительные параметры
            col_params1, col_params2 = st.columns(2)
            
            with col_params1:
                doc_style = st.selectbox(
                    "Стиль документа:",
                    ["Формальный", "Стандартный", "Упрощенный", "Детализированный"]
                )
                
                doc_party = st.selectbox(
                    "Чья позиция:",
                    ["Автора документа", "Принимающей стороны", "Нейтральная"]
                )
            
            with col_params2:
                include_comments = st.checkbox("Добавить комментарии к пунктам", value=True)
                include_alternatives = st.checkbox("Включить альтернативные варианты", value=False)
            
            if st.button("🔄 СГЕНЕРИРОВАТЬ ДОКУМЕНТ", use_container_width=True):
                with st.spinner("Создаю документ..."):
                    prompt = f"""
                    Ты - юрист, создающий документ с нуля.
                    
                    ЗАПРОС ПОЛЬЗОВАТЕЛЯ:
                    {doc_description}
                    
                    ТРЕБОВАНИЯ:
                    1. Создай полноценный юридический документ
                    2. Стиль: {doc_style}
                    3. Позиция: {doc_party}
                    4. Включи все необходимые разделы
                    5. Учитывай законодательство {loc}
                    6. Документ должен быть готов к использованию
                    
                    {"7. Добавь комментарии к сложным пунктам" if include_comments else ""}
                    {"8. Предложи альтернативные формулировки для ключевых условий" if include_alternatives else ""}
                    
                    СТРУКТУРА ДОКУМЕНТА:
                    - Преамбула (реквизиты сторон)
                    - Предмет договора
                    - Права и обязанности сторон
                    - Сроки и условия
                    - Оплата и расчеты
                    - Ответственность сторон
                    - Форс-мажор
                    - Разрешение споров
                    - Заключительные положения
                    - Реквизиты и подписи
                    
                    Сделай документ профессиональным и юридически корректным.
                    """
                    
                    generated_doc = call_gemini_safe(prompt, "")
                    
                    if generated_doc:
                        st.markdown("### 📄 Сгенерированный документ")
                        st.markdown(generated_doc)
                        
                        # Определяем название документа из первого заголовка
                        doc_title = "Сгенерированный документ"
                        lines = generated_doc.split('\n')
                        for line in lines:
                            if line.startswith('# ') and len(line) > 2:
                                doc_title = line[2:].strip()
                                break
                        
                        # Кнопка скачивания
                        docx_file = create_docx(generated_doc, doc_title)
                        st.download_button(
                            "📥 Скачать документ",
                            data=docx_file,
                            file_name=f"{doc_title.replace(' ', '_')}.docx",
                            use_container_width=True
                        )

# -------------------
# 7. Футер
# -------------------
st.markdown("---")
footer_col1, footer_col2, footer_col3 = st.columns(3)

with footer_col1:
    st.caption("© 2024 LegalAI Enterprise Pro")
    st.caption("Версия 2.0.0")

with footer_col2:
    st.caption("Powered by Google Gemini AI")
    st.caption("Для образовательных целей")

with footer_col3:
    st.caption("🔒 Ваши данные обрабатываются безопасно")
    st.caption("Поддержка: support@legalai.pro")

# Скрыть Streamlit элементы по умолчанию
hide_streamlit_style = """
<style>
#MainMenu {visibility: hidden;}
footer {visibility: hidden;}
header {visibility: hidden;}
</style>
"""
st.markdown(hide_streamlit_style, unsafe_allow_html=True)
